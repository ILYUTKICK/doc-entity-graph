#!/usr/bin/env python3
"""Create a small DOCX input for demonstrating the document graph pipeline."""

from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
RAW_DEMO = ROOT / "data" / "raw_teacher_demo"
DOCX_PATH = RAW_DEMO / "demo_teacher_pipeline.docx"
TXT_PATH = RAW_DEMO / "demo_teacher_pipeline_source.txt"


SOURCE_TEXT = """\
Демо-документ для проверки графа сущностей

Введение

Этот документ создан специально для демонстрации пайплайна doc-entity-graph. В нём повторяются ключевые сущности: Bank of Russia, Moscow Exchange, Sberbank, Gazprom, USD/RUB, inflation, key rate, ARIMA, ADF test и Ljung-Box test.

Сценарий документа описывает анализ макроэкономических данных за 2024-2026 годы. Аналитик сравнивает inflation, key rate и USD/RUB, чтобы объяснить, как решения Bank of Russia влияют на рынок и прогнозы компаний Sberbank и Gazprom.

Таблица 1. Исходные показатели

Период | Inflation | Key rate | USD/RUB | Комментарий
2024 Q1 | 7.4 | 16.0 | 91.2 | Bank of Russia сохраняет жёсткую политику
2024 Q2 | 8.1 | 16.0 | 88.7 | Moscow Exchange показывает укрепление рубля
2024 Q3 | 8.9 | 18.0 | 92.5 | Sberbank повышает прогноз инфляции
2024 Q4 | 9.2 | 21.0 | 97.8 | Gazprom учитывает валютные риски

Рисунок 1. Динамика inflation и key rate

После роста inflation Bank of Russia повышает key rate. На графике видно, что key rate реагирует на инфляционный шок с задержкой. Такая связь важна для ARIMA-модели, потому что ряд key rate меняется дискретно, а inflation меняется более плавно.

Моделирование

Для inflation используется ARIMA(1,0,1), для key rate используется ARIMA(1,1,0), для USD/RUB используется ARIMA(0,1,1). Перед построением моделей применяется ADF test. После оценки остатков применяется Ljung-Box test, чтобы проверить автокорреляцию.

Таблица 2. Выбор моделей

Ряд | Модель | ADF p-value | Ljung-Box p-value | Вывод
Inflation | ARIMA(1,0,1) | 0.031 | 0.42 | Остатки похожи на белый шум
Key rate | ARIMA(1,1,0) | 0.018 | 0.37 | Модель учитывает смену политики Bank of Russia
USD/RUB | ARIMA(0,1,1) | 0.011 | 0.55 | Валютный ряд требует дифференцирования

Рисунок 2. Прогноз USD/RUB на 2026 год

Прогноз показывает, что USD/RUB может вырасти при сохранении высокой inflation. Moscow Exchange и Sberbank используют похожий сценарий в стресс-тестах. Gazprom чувствителен к курсу USD/RUB, потому что валютная выручка влияет на финансовый результат.

Итог

Документ нужен для проверки того, что проект связывает сущности с чанками, таблицами, рисунками и подписями. В linking-графе должны появиться связи между Bank of Russia, inflation, key rate, USD/RUB, ARIMA, ADF test, Ljung-Box test, таблицами и рисунками.
"""


def add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def create_line_chart(path: Path, title: str, labels: list[str], series: list[tuple[str, list[float], str]]) -> None:
    width, height = 900, 430
    margin_left, margin_right = 80, 40
    margin_top, margin_bottom = 70, 80
    plot_w = width - margin_left - margin_right
    plot_h = height - margin_top - margin_bottom

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default()

    all_values = [value for _, values, _ in series for value in values]
    min_v = min(all_values) - 2
    max_v = max(all_values) + 2

    def x_pos(i: int) -> float:
        if len(labels) == 1:
            return margin_left + plot_w / 2
        return margin_left + i * plot_w / (len(labels) - 1)

    def y_pos(value: float) -> float:
        return margin_top + (max_v - value) * plot_h / (max_v - min_v)

    draw.text((margin_left, 25), title, fill="#111827", font=font)
    draw.line((margin_left, margin_top, margin_left, margin_top + plot_h), fill="#374151", width=2)
    draw.line((margin_left, margin_top + plot_h, margin_left + plot_w, margin_top + plot_h), fill="#374151", width=2)

    for step in range(5):
        value = min_v + step * (max_v - min_v) / 4
        y = y_pos(value)
        draw.line((margin_left, y, margin_left + plot_w, y), fill="#e5e7eb", width=1)
        draw.text((20, y - 7), f"{value:.1f}", fill="#4b5563", font=font)

    for i, label in enumerate(labels):
        x = x_pos(i)
        draw.line((x, margin_top + plot_h, x, margin_top + plot_h + 5), fill="#374151", width=1)
        draw.text((x - 24, margin_top + plot_h + 15), label, fill="#4b5563", font=font)

    legend_x = margin_left + 10
    legend_y = height - 35
    for name, values, color in series:
        points = [(x_pos(i), y_pos(value)) for i, value in enumerate(values)]
        draw.line(points, fill=color, width=4)
        for x, y in points:
            draw.ellipse((x - 5, y - 5, x + 5, y + 5), fill=color, outline=color)
        draw.rectangle((legend_x, legend_y, legend_x + 18, legend_y + 10), fill=color)
        draw.text((legend_x + 24, legend_y - 2), name, fill="#111827", font=font)
        legend_x += 170

    image.save(path)


def build_docx() -> None:
    RAW_DEMO.mkdir(parents=True, exist_ok=True)
    TXT_PATH.write_text(SOURCE_TEXT, encoding="utf-8")

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    document.add_heading("Демо-документ для проверки графа сущностей", level=0)
    document.add_paragraph(
        "Документ создан специально для запуска пайплайна doc-entity-graph. "
        "В нём есть повторяющиеся сущности, таблицы, графики и подписи."
    )

    document.add_heading("1. Введение", level=1)
    add_paragraph(
        document,
        "В документе повторяются ключевые сущности: Bank of Russia, Moscow Exchange, "
        "Sberbank, Gazprom, USD/RUB, inflation, key rate, ARIMA, ADF test и Ljung-Box test.",
    )
    add_paragraph(
        document,
        "Сценарий описывает анализ макроэкономических данных за 2024-2026 годы. "
        "Аналитик сравнивает inflation, key rate и USD/RUB, чтобы объяснить, как решения "
        "Bank of Russia влияют на рынок и прогнозы компаний Sberbank и Gazprom.",
    )

    add_caption(document, "Таблица 1. Исходные показатели")
    add_table(
        document,
        ["Период", "Inflation", "Key rate", "USD/RUB", "Комментарий"],
        [
            ["2024 Q1", "7.4", "16.0", "91.2", "Bank of Russia сохраняет жёсткую политику"],
            ["2024 Q2", "8.1", "16.0", "88.7", "Moscow Exchange показывает укрепление рубля"],
            ["2024 Q3", "8.9", "18.0", "92.5", "Sberbank повышает прогноз инфляции"],
            ["2024 Q4", "9.2", "21.0", "97.8", "Gazprom учитывает валютные риски"],
        ],
    )

    with TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        fig1 = temp_path / "demo_fig1.png"
        fig2 = temp_path / "demo_fig2.png"
        create_line_chart(
            fig1,
            "Inflation and key rate, 2024",
            ["Q1", "Q2", "Q3", "Q4"],
            [
                ("Inflation", [7.4, 8.1, 8.9, 9.2], "#d85a30"),
                ("Key rate", [16.0, 16.0, 18.0, 21.0], "#1d9e75"),
            ],
        )
        document.add_picture(str(fig1), width=Inches(5.9))
        add_caption(document, "Рисунок 1. Динамика inflation и key rate")
        add_paragraph(
            document,
            "После роста inflation Bank of Russia повышает key rate. На графике видно, что "
            "key rate реагирует на инфляционный шок с задержкой. Такая связь важна для ARIMA-модели.",
        )

        document.add_heading("2. Моделирование", level=1)
        add_paragraph(
            document,
            "Для inflation используется ARIMA(1,0,1), для key rate используется ARIMA(1,1,0), "
            "для USD/RUB используется ARIMA(0,1,1). Перед построением моделей применяется "
            "ADF test. После оценки остатков применяется Ljung-Box test.",
        )

        add_caption(document, "Таблица 2. Выбор моделей")
        add_table(
            document,
            ["Ряд", "Модель", "ADF p-value", "Ljung-Box p-value", "Вывод"],
            [
                ["Inflation", "ARIMA(1,0,1)", "0.031", "0.42", "Остатки похожи на белый шум"],
                ["Key rate", "ARIMA(1,1,0)", "0.018", "0.37", "Модель учитывает Bank of Russia"],
                ["USD/RUB", "ARIMA(0,1,1)", "0.011", "0.55", "Ряд требует дифференцирования"],
            ],
        )

        create_line_chart(
            fig2,
            "USD/RUB forecast, 2026",
            ["Jan", "Feb", "Mar", "Apr"],
            [
                ("Actual", [97.8, 98.5, 99.1, 100.3], "#378add"),
                ("Forecast", [98.0, 99.2, 101.0, 102.4], "#7f77dd"),
            ],
        )
        document.add_picture(str(fig2), width=Inches(5.9))
        add_caption(document, "Рисунок 2. Прогноз USD/RUB на 2026 год")
        add_paragraph(
            document,
            "Прогноз показывает, что USD/RUB может вырасти при сохранении высокой inflation. "
            "Moscow Exchange и Sberbank используют похожий сценарий в стресс-тестах. "
            "Gazprom чувствителен к курсу USD/RUB.",
        )

    document.add_heading("3. Итог", level=1)
    add_paragraph(
        document,
        "Документ нужен для проверки того, что проект связывает сущности с чанками, таблицами, "
        "рисунками и подписями. В linking-графе должны появиться связи между Bank of Russia, "
        "inflation, key rate, USD/RUB, ARIMA, ADF test, Ljung-Box test, таблицами и рисунками.",
    )

    document.save(DOCX_PATH)


def main() -> None:
    build_docx()
    print(f"Created DOCX: {DOCX_PATH}")
    print(f"Created source text: {TXT_PATH}")


if __name__ == "__main__":
    main()
